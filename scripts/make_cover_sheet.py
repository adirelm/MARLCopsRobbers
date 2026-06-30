"""Generate the Moodle cover sheet adrl-001-ex06.pdf from the Word template.

Reads PII (name/id/github) from the GIT-IGNORED ``players.local.yaml`` and the two
human-owned fields it does NOT contain — ``self_score`` (§1.4 non-delegable) and
``name_he: {first, last}`` (Hebrew name) — if present there; otherwise leaves them
blank for the human to fill. Fills the provided template (fields only, no extra text;
ex06 §"Standing rules") and converts to PDF via LibreOffice (``soffice``). Output PDF +
the intermediate .docx are git-ignored (Moodle-only). NO PII is printed.

Run:  uv run --with python-docx --with pyyaml python scripts/make_cover_sheet.py
"""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

import docx
import yaml

_ROOT = Path(__file__).resolve().parents[1]
_TEMPLATE = _ROOT / "instructions" / "assignment-6" / "biu-rl07-ex01-template.docx"
_SOFFICE_CANDIDATES = (
    "soffice",
    "libreoffice",
    "/opt/homebrew/bin/soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
)


def _resolve_soffice() -> str:
    """Return the first available LibreOffice ``soffice`` binary (for docx -> pdf)."""
    for cand in _SOFFICE_CANDIDATES:
        found = shutil.which(cand) or (cand if Path(cand).exists() else None)
        if found:
            return found
    raise RuntimeError("LibreOffice 'soffice' not found (needed for docx -> pdf conversion)")


def _append(paragraph: object, value: str) -> None:
    """Append ``value`` after a label paragraph's text (fields-only fill)."""
    if not value:
        return
    if paragraph.runs:
        paragraph.runs[0].text = paragraph.runs[0].text.rstrip() + "   " + value
    else:
        paragraph.add_run("   " + value)


def main() -> Path:
    """Fill the template from players.local.yaml + convert to a git-ignored PDF."""
    players = yaml.safe_load((_ROOT / "players.local.yaml").read_text(encoding="utf-8"))
    student = players["students"][0]
    name_parts = str(student["full_name"]).split()
    first_en = name_parts[0] if name_parts else ""
    last_en = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
    name_he = players.get("name_he") or {}
    score = str(players.get("self_score", "")).strip()  # §1.4: human-owned, blank if absent

    doc = docx.Document(str(_TEMPLATE))
    p = doc.paragraphs
    _append(p[0], "06")  # exercise number
    _append(p[2], str(players["group_name"]))  # group ID code
    _append(p[4], score)  # self-score recommendation (human)
    _append(p[7], str(student["id"]))  # Student 1 ID
    _append(p[8], first_en)  # Student 1 first (EN)
    _append(p[9], last_en)  # Student 1 last (EN)
    _append(p[10], str(name_he.get("first", "")))  # Student 1 first (HE)
    _append(p[11], str(name_he.get("last", "")))  # Student 1 last (HE)
    # Student 2 (paras 13-18) left blank — SOLO submission.
    _append(p[20], str(players["github_repo"]))  # GitHub link
    _append(p[21], "no")  # late submission?

    out_docx = _ROOT / "adrl-001-ex06.docx"
    doc.save(str(out_docx))
    subprocess.run(
        [_resolve_soffice(), "--headless", "--convert-to", "pdf", "--outdir", str(_ROOT), str(out_docx)],
        check=True,
        capture_output=True,
    )
    pdf = _ROOT / "adrl-001-ex06.pdf"
    missing = [k for k, v in {"self_score": score, "name_he": name_he}.items() if not v]
    print(f"wrote {pdf.name} (solo); blank human-owned fields: {missing or 'none'}")
    return pdf


if __name__ == "__main__":
    main()
